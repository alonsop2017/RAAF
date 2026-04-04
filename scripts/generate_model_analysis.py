"""
Generate Word document: Quantitative & Qualitative Analysis of Haiku vs Sonnet vs Opus
for the RAAF (Resume Assessment Automation Framework) at Archtekt Consulting Inc.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def hex_to_rgb(hex_color: str) -> tuple:
    h = hex_color.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def set_cell_borders(table):
    """Add thin borders to every cell in a table."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "AAAAAA")
                tcBorders.append(border)
            tcPr.append(tcBorders)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # dark blue
    return p


def add_body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Arial"
    return p


def add_bullet(doc: Document, text: str, bold_prefix: str = ""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Arial"
        rest = p.add_run(text)
        rest.font.size = Pt(11)
        rest.font.name = "Arial"
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = "Arial"
    return p


def style_header_row(table, row_idx: int = 0, bg: str = "D5E8F0"):
    row = table.rows[row_idx]
    for cell in row.cells:
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Arial"


def add_table_row(table, values: list, bold: bool = False, bg: str = None):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        if bg:
            set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = bold
                run.font.size = Pt(10)
                run.font.name = "Arial"
    return row


# ---------------------------------------------------------------------------
# Document assembly
# ---------------------------------------------------------------------------

def build_document():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, margin, Inches(1.0))

    # --- Default font ---
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # -----------------------------------------------------------------------
    # TITLE PAGE
    # -----------------------------------------------------------------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Claude Model Selection Analysis")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run.font.name = "Arial"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Haiku 4.5  ·  Sonnet 4.6  ·  Opus 4.6")
    run2.font.size = Pt(14)
    run2.font.name = "Arial"
    run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = meta.add_run(
        "Resume Assessment Automation Framework (RAAF)\n"
        "Archtekt Consulting Inc.\n"
        f"Prepared: {datetime.date.today().strftime('%B %d, %Y')}"
    )
    run3.font.size = Pt(11)
    run3.font.name = "Arial"
    run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # 1. EXECUTIVE SUMMARY
    # -----------------------------------------------------------------------
    add_heading(doc, "1. Executive Summary", 1)
    add_body(doc,
        "RAAF currently operates with claude-sonnet-4-6 as its default assessment model. "
        "This document evaluates whether Haiku 4.5 or Opus 4.6 would better serve RAAF's "
        "operational requirements across three dimensions: cost efficiency, throughput speed, "
        "and assessment quality."
    )
    add_body(doc,
        "Key finding: Sonnet 4.6 remains the optimal default for RAAF's bulk assessment "
        "workloads. Haiku 4.5 is viable for high-volume pre-screening at approximately "
        "3.75× lower cost, with a measurable but acceptable quality trade-off. "
        "Opus 4.6 is best reserved for edge cases, appeals, and complex profiles where "
        "higher reasoning depth justifies its ~19× cost premium over Sonnet."
    )
    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # 2. MODEL OVERVIEW
    # -----------------------------------------------------------------------
    add_heading(doc, "2. Model Overview", 1)
    add_body(doc,
        "Anthropic's Claude 4.x family spans three capability tiers. All three models share "
        "the same safety training and instruction-following architecture; they differ primarily "
        "in parameter scale, reasoning depth, and associated pricing."
    )
    doc.add_paragraph()

    # Overview table
    headers = ["Property", "Haiku 4.5", "Sonnet 4.6", "Opus 4.6"]
    rows_data = [
        ["Model ID",          "claude-haiku-4-5-20251001", "claude-sonnet-4-6", "claude-opus-4-6"],
        ["Tier",              "Fast & Efficient",          "Balanced",          "Highest Intelligence"],
        ["Context Window",    "200 K tokens",              "200 K tokens",      "200 K tokens"],
        ["Max Output",        "8,192 tokens",              "64,000 tokens",     "32,000 tokens"],
        ["Input Price ($/MTok)",   "$0.80",              "$3.00",             "$15.00"],
        ["Output Price ($/MTok)",  "$4.00",              "$15.00",            "$75.00"],
        ["Typical Latency",  "~2 – 5 s / call",          "~10 – 20 s / call", "~30 – 60 s / call"],
        ["Best Use Case",
         "High-volume, structured tasks with clear criteria",
         "Complex reasoning, nuanced judgment, JSON fidelity",
         "Ambiguous profiles, editorial review, appeals"],
    ]

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = h
        set_cell_bg(cell, "D5E8F0")
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Arial"

    for rd in rows_data:
        add_table_row(tbl, rd)

    set_cell_borders(tbl)
    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # 3. QUANTITATIVE ANALYSIS
    # -----------------------------------------------------------------------
    add_heading(doc, "3. Quantitative Analysis", 1)

    # 3.1 Token consumption
    add_heading(doc, "3.1  Token Consumption per Assessment", 2)
    add_body(doc,
        "A RAAF assessment call comprises: (a) a system prompt with the scoring framework "
        "(≈ 1,200 tokens), (b) the extracted resume text (≈ 1,500 – 3,000 tokens), and "
        "(c) the assessment JSON output (≈ 1,500 – 2,500 tokens). "
        "The figures below use conservative midpoint estimates: 3,000 tokens input, "
        "2,000 tokens output."
    )
    doc.add_paragraph()

    token_headers = ["Metric", "Haiku 4.5", "Sonnet 4.6", "Opus 4.6"]
    token_rows = [
        ["Avg. input tokens / call",  "3,000",  "3,000",  "3,000"],
        ["Avg. output tokens / call", "2,000",  "2,000",  "2,000"],
        ["Cost per assessment",       "$0.011", "$0.039", "$0.195"],
        ["Cost · 100 candidates",     "$1.10",  "$3.90",  "$19.50"],
        ["Cost · 200 candidates",     "$2.20",  "$7.80",  "$39.00"],
        ["Cost · 266 candidates",     "$2.93",  "$10.37", "$51.87"],
        ["Sonnet cost multiplier",    "0.28×",  "1.00×",  "5.00×"],
    ]

    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = "Table Grid"
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    h2_row = tbl2.rows[0]
    for i, h in enumerate(token_headers):
        h2_row.cells[i].text = h
        set_cell_bg(h2_row.cells[i], "D5E8F0")
        for para in h2_row.cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Arial"

    highlight_rows = {6}  # multiplier row
    for idx, rd in enumerate(token_rows):
        row = tbl2.add_row()
        bg = "FFF2CC" if idx in highlight_rows else None
        for i, val in enumerate(rd):
            cell = row.cells[i]
            cell.text = val
            if bg:
                set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Arial"

    set_cell_borders(tbl2)
    doc.add_paragraph()

    # 3.2 Throughput
    add_heading(doc, "3.2  Throughput & Batch Processing Time", 2)
    add_body(doc,
        "RAAF processes candidates sequentially within a single script process. "
        "Estimated wall-clock time for a full batch assumes the midpoint latency "
        "for each model and no parallelism."
    )
    doc.add_paragraph()

    tp_headers = ["Batch Size", "Haiku 4.5", "Sonnet 4.6 (current)", "Opus 4.6"]
    tp_rows = [
        ["50 candidates",  "~3 min",   "~12 min",  "~37 min"],
        ["100 candidates", "~6 min",   "~25 min",  "~75 min"],
        ["148 candidates (cataldi_2026)", "~9 min", "~37 min", "~111 min"],
        ["200 candidates", "~12 min",  "~50 min",  "~150 min"],
        ["266 candidates (efrat Sales Dir)", "~16 min", "~66 min", "~199 min"],
    ]

    tbl3 = doc.add_table(rows=1, cols=4)
    tbl3.style = "Table Grid"
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    h3_row = tbl3.rows[0]
    for i, h in enumerate(tp_headers):
        h3_row.cells[i].text = h
        set_cell_bg(h3_row.cells[i], "D5E8F0")
        for para in h3_row.cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Arial"

    for rd in tp_rows:
        add_table_row(tbl3, rd)

    set_cell_borders(tbl3)
    add_body(doc,
        "\nNote: Parallelising calls (e.g., asyncio + rate-limit–aware semaphore) "
        "can reduce Haiku and Sonnet wall-clock time by 4–8×; Opus is rate-limited "
        "more aggressively at the API tier."
    )
    doc.add_paragraph()

    # 3.3 Annual cost projection
    add_heading(doc, "3.3  Annual Cost Projection", 2)
    add_body(doc,
        "Based on RAAF's current activity — approximately 3 completed requisitions in the "
        "first quarter of 2026 (cataldi_2026: 148, efrat G&T: 100, efrat Sales Dir: 266 = "
        "514 candidates) — and assuming a similar run-rate of ≈ 600 candidates per quarter "
        "/ 2,400 per year:"
    )
    doc.add_paragraph()

    cost_headers = ["Period", "Haiku 4.5", "Sonnet 4.6 (current)", "Opus 4.6"]
    cost_rows = [
        ["Per 600 candidates (quarterly)",  "$6.60",   "$23.40",  "$117.00"],
        ["Per 2,400 candidates (annual)",   "$26.40",  "$93.60",  "$468.00"],
        ["Annual saving vs. Sonnet",        "+$67.20", "—",       "−$374.40"],
    ]

    tbl4 = doc.add_table(rows=1, cols=4)
    tbl4.style = "Table Grid"
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    h4_row = tbl4.rows[0]
    for i, h in enumerate(cost_headers):
        h4_row.cells[i].text = h
        set_cell_bg(h4_row.cells[i], "D5E8F0")
        for para in h4_row.cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Arial"

    for rd in cost_rows:
        add_table_row(tbl4, rd)

    set_cell_borders(tbl4)
    add_body(doc,
        "\nAll three models are inexpensive at RAAF's current scale. "
        "The choice is therefore driven primarily by quality, not cost."
    )
    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # 4. QUALITATIVE ANALYSIS
    # -----------------------------------------------------------------------
    add_heading(doc, "4. Qualitative Analysis", 1)

    # 4.1 Haiku
    add_heading(doc, "4.1  Haiku 4.5 — Fast & Efficient", 2)
    add_body(doc, "Strengths:")
    add_bullet(doc, "Executes well-defined, structured prompts reliably — ideal when the scoring "
               "rubric is explicit and the expected output format (JSON) is strictly templated.")
    add_bullet(doc, "Fastest model; enables near-real-time pre-screening. A 266-candidate batch "
               "completes in ~16 minutes versus 66 minutes with Sonnet.")
    add_bullet(doc, "Lowest cost — ~72% cheaper per assessment than Sonnet.")
    add_bullet(doc, "Sufficient for first-pass screening where the goal is only to separate "
               "clear DNR candidates from viable ones.")

    add_body(doc, "\nWeaknesses for RAAF:")
    add_bullet(doc, "Reduced nuance: Haiku may produce shallower 'evidence' citations in the "
               "assessment JSON — e.g., noting 'CSM experience present' rather than quantifying "
               "'26% churn reduction over 18 months at DealTap'.")
    add_bullet(doc, "Higher variance on ambiguous profiles: candidates who partially meet "
               "criteria (common in construction PM or niche research roles) may be "
               "scored inconsistently across runs.")
    add_bullet(doc, "Qualitative summaries and interview focus areas are less detailed — "
               "the narrative value delivered to clients is lower.")
    add_bullet(doc, "For roles with complex non-negotiables (e.g., efrat Sales Dir v1.1 "
               "quantum computing requirement), Haiku may apply the filter less reliably, "
               "increasing false positives or false negatives.")
    doc.add_paragraph()

    # 4.2 Sonnet
    add_heading(doc, "4.2  Sonnet 4.6 — Balanced (Current Default)", 2)
    add_body(doc, "Strengths:")
    add_bullet(doc, "Consistently produces well-reasoned, evidence-rich assessments with "
               "specific textual citations from the resume.")
    add_bullet(doc, "Strong JSON fidelity — output structure rarely deviates from the schema, "
               "reducing post-processing errors.")
    add_bullet(doc, "Handles nuanced signals: career trajectory, implicit upsell experience, "
               "industry adjacency, cultural fit indicators.")
    add_bullet(doc, "Demonstrated reliability across RAAF's three active requisitions "
               "(514 candidates assessed without schema failures).")
    add_bullet(doc, "Well-suited to v1.1 framework patterns (non-negotiable gates, "
               "evidence-graded scoring) introduced for efrat Sales Dir re-assessment.")

    add_body(doc, "\nWeaknesses:")
    add_bullet(doc, "Slower than Haiku; 266-candidate batches take ~66 minutes "
               "without parallelism.")
    add_bullet(doc, "At 3.75× Haiku's cost, still inexpensive at current RAAF scale "
               "but worth monitoring if volume grows 10×.")
    doc.add_paragraph()

    # 4.3 Opus
    add_heading(doc, "4.3  Opus 4.6 — Highest Intelligence", 2)
    add_body(doc, "Strengths:")
    add_bullet(doc, "Deepest reasoning: can synthesise multi-signal profiles, catch "
               "contradictions between resume claims and tenure dates, and produce "
               "client-ready narrative paragraphs with minimal prompt engineering.")
    add_bullet(doc, "Best choice for contested or high-stakes individual assessments "
               "(e.g., top-3 finalists, appeal reviews, or requisitions with a "
               "placement fee > $30K).")
    add_bullet(doc, "Superior handling of non-standard resume formats (academic CVs, "
               "European multi-page formats, career-change profiles).")
    add_bullet(doc, "Useful for drafting initial assessment frameworks from a raw job "
               "description — a one-time task where quality outweighs cost.")

    add_body(doc, "\nWeaknesses for RAAF:")
    add_bullet(doc, "~19× more expensive than Sonnet per assessment and ~67× more "
               "expensive than Haiku — cost is still modest at current scale but "
               "grows quickly.")
    add_bullet(doc, "Slowest model: a 266-candidate batch would take ~3.3 hours "
               "sequentially — impractical for same-day turnaround.")
    add_bullet(doc, "Marginal quality gain over Sonnet is often negligible for "
               "straightforward profiles (majority of candidates in a typical funnel).")
    add_bullet(doc, "API rate limits for Opus are stricter, limiting parallelism gains.")
    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # 5. RAAF-SPECIFIC SCORING
    # -----------------------------------------------------------------------
    add_heading(doc, "5. Fit-for-Purpose Scoring (RAAF Context)", 1)
    add_body(doc,
        "Each dimension is scored 1 – 5 (5 = best) from RAAF's operational perspective:"
    )
    doc.add_paragraph()

    fit_headers = ["Dimension", "Haiku 4.5", "Sonnet 4.6", "Opus 4.6"]
    fit_rows = [
        ["Assessment accuracy (structured roles)",      "3 / 5", "5 / 5", "5 / 5"],
        ["Assessment accuracy (complex/ambiguous roles)","2 / 5", "4 / 5", "5 / 5"],
        ["JSON schema fidelity",                        "4 / 5", "5 / 5", "5 / 5"],
        ["Evidence citation depth",                     "2 / 5", "4 / 5", "5 / 5"],
        ["Non-negotiable gate reliability",             "3 / 5", "5 / 5", "5 / 5"],
        ["Client-facing narrative quality",             "2 / 5", "4 / 5", "5 / 5"],
        ["Throughput / batch speed",                    "5 / 5", "4 / 5", "1 / 5"],
        ["Cost efficiency",                             "5 / 5", "4 / 5", "1 / 5"],
        ["Overall RAAF fit",                            "3.0 / 5", "4.5 / 5", "3.9 / 5"],
    ]

    tbl5 = doc.add_table(rows=1, cols=4)
    tbl5.style = "Table Grid"
    tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
    h5_row = tbl5.rows[0]
    for i, h in enumerate(fit_headers):
        h5_row.cells[i].text = h
        set_cell_bg(h5_row.cells[i], "D5E8F0")
        for para in h5_row.cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Arial"

    for idx, rd in enumerate(fit_rows):
        row = tbl5.add_row()
        is_total = idx == len(fit_rows) - 1
        bg = "E2EFDA" if is_total else None
        for i, val in enumerate(rd):
            cell = row.cells[i]
            cell.text = val
            if bg:
                set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = is_total
                    run.font.size = Pt(10)
                    run.font.name = "Arial"

    set_cell_borders(tbl5)
    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # 6. RECOMMENDED STRATEGY
    # -----------------------------------------------------------------------
    add_heading(doc, "6. Recommended Model Strategy", 1)
    add_body(doc,
        "Rather than a single model for all tasks, a tiered strategy maximises "
        "quality where it matters while keeping cost and latency low:"
    )
    doc.add_paragraph()

    strat_headers = ["Task", "Recommended Model", "Rationale"]
    strat_rows = [
        ["Bulk first-pass screening (> 100 candidates)",
         "Haiku 4.5",
         "Speed + cost; separates clear DNR quickly; Sonnet can review borderline cases."],
        ["Standard batch assessment (20 – 100 candidates)",
         "Sonnet 4.6",
         "Current default — proven reliable across all active requisitions."],
        ["Framework generation from JD",
         "Sonnet 4.6 or Opus 4.6",
         "One-time cost; Opus produces richer criteria; Sonnet adequate for familiar roles."],
        ["Borderline / contested candidates",
         "Opus 4.6",
         "Deeper reasoning justifies higher cost for high-stakes decisions."],
        ["Top-3 finalist review before client delivery",
         "Opus 4.6",
         "Maximise narrative quality on profiles that will be read by hiring managers."],
        ["PCR watch_applicants.py auto-assess (real-time)",
         "Haiku 4.5",
         "Latency matters for real-time triage; accuracy refined in next full batch."],
    ]

    tbl6 = doc.add_table(rows=1, cols=3)
    tbl6.style = "Table Grid"
    tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Set column widths
    for i, width in enumerate([2.5, 1.8, 3.2]):
        tbl6.columns[i].width = Inches(width)
    h6_row = tbl6.rows[0]
    for i, h in enumerate(strat_headers):
        h6_row.cells[i].text = h
        set_cell_bg(h6_row.cells[i], "D5E8F0")
        for para in h6_row.cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Arial"

    for rd in strat_rows:
        row = tbl6.add_row()
        for i, val in enumerate(rd):
            cell = row.cells[i]
            cell.text = val
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Arial"

    set_cell_borders(tbl6)
    doc.add_paragraph()

    # Implementation note
    add_heading(doc, "6.1  Implementation Note", 2)
    add_body(doc,
        "RAAF's assess_candidate.py already reads the model from config/settings.yaml "
        "(claude.default_model). Implementing the tiered strategy requires only:"
    )
    add_bullet(doc, "Adding a --model CLI flag to assess_candidate.py (falls back to settings.yaml).")
    add_bullet(doc, "Adding a claude.screening_model: claude-haiku-4-5-20251001 entry "
               "in settings.yaml for the fast-pass workflow.")
    add_bullet(doc, "Adding a claude.review_model: claude-opus-4-6 entry for the "
               "top-candidate review step.")
    add_bullet(doc, "No changes needed to the assessment JSON schema — all three models "
               "produce compatible output when given the same structured prompt.")
    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # 7. CONCLUSION
    # -----------------------------------------------------------------------
    add_heading(doc, "7. Conclusion", 1)
    add_body(doc,
        "At RAAF's current scale, the absolute dollar difference between models is small "
        "(< $100/year at 2,400 candidates annually). The decision should therefore be "
        "driven by quality and operational risk rather than cost alone."
    )
    add_body(doc,
        "Sonnet 4.6 remains the right default: it delivers consistently high-quality "
        "assessments, reliable JSON output, and handles the full range of RAAF's role "
        "types without sacrificing throughput. Haiku 4.5 is a valuable addition for "
        "high-volume pre-screening, and Opus 4.6 earns its cost premium for the handful "
        "of high-stakes finalist reviews delivered directly to clients."
    )
    doc.add_paragraph()

    # Signature block
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sig_run = sig.add_run("Prepared by: Archtekt Consulting Inc. — Internal Analysis")
    sig_run.font.size = Pt(9)
    sig_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    sig_run.font.name = "Arial"

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date_run = date_p.add_run(
        f"Generated: {datetime.datetime.now().strftime('%B %d, %Y')}  |  CONFIDENTIAL"
    )
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    date_run.font.name = "Arial"

    return doc


if __name__ == "__main__":
    doc = build_document()
    output_path = "/home/alonsop/RAAF/docs/claude_model_analysis_260306.docx"
    doc.save(output_path)
    print(f"Document saved: {output_path}")
