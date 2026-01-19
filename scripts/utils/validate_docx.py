#!/usr/bin/env python3
"""
Validate DOCX file integrity.
Checks that a DOCX file is valid and readable.
"""

import sys
import zipfile
from pathlib import Path

try:
    from docx import Document
except ImportError:
    Document = None


def validate_docx(docx_path: str | Path) -> dict:
    """
    Validate a DOCX file.

    Args:
        docx_path: Path to DOCX file

    Returns:
        Validation results
    """
    docx_path = Path(docx_path)

    results = {
        "valid": True,
        "file": str(docx_path),
        "errors": [],
        "warnings": [],
        "info": {}
    }

    # Check file exists
    if not docx_path.exists():
        results["valid"] = False
        results["errors"].append("File not found")
        return results

    # Check extension
    if docx_path.suffix.lower() != ".docx":
        results["warnings"].append(f"Unexpected extension: {docx_path.suffix}")

    # Check it's a valid ZIP (DOCX is a ZIP archive)
    try:
        with zipfile.ZipFile(docx_path, 'r') as zf:
            # Check for required DOCX components
            namelist = zf.namelist()

            required = [
                "[Content_Types].xml",
                "word/document.xml"
            ]

            for req in required:
                if req not in namelist:
                    results["valid"] = False
                    results["errors"].append(f"Missing required component: {req}")

            results["info"]["components"] = len(namelist)

    except zipfile.BadZipFile:
        results["valid"] = False
        results["errors"].append("Not a valid ZIP/DOCX file")
        return results

    # Try to open with python-docx
    if Document:
        try:
            doc = Document(docx_path)

            results["info"]["paragraphs"] = len(doc.paragraphs)
            results["info"]["tables"] = len(doc.tables)
            results["info"]["sections"] = len(doc.sections)

            # Check for content
            if len(doc.paragraphs) == 0:
                results["warnings"].append("Document has no paragraphs")

        except Exception as e:
            results["valid"] = False
            results["errors"].append(f"python-docx error: {e}")
    else:
        results["warnings"].append("python-docx not installed - skipping deep validation")

    # File size check
    size_mb = docx_path.stat().st_size / (1024 * 1024)
    results["info"]["size_mb"] = round(size_mb, 2)

    if size_mb > 10:
        results["warnings"].append(f"Large file: {size_mb:.1f} MB")

    return results


def main():
    import argparse
    import json

    parser = argparse.ArgumentParser(description="Validate DOCX file")
    parser.add_argument("file", help="DOCX file to validate")
    parser.add_argument("--json", action="store_true", help="Output as JSON")
    args = parser.parse_args()

    results = validate_docx(args.file)

    if args.json:
        print(json.dumps(results, indent=2))
    else:
        status = "✓ VALID" if results["valid"] else "✗ INVALID"
        print(f"\nDOCX Validation: {results['file']}")
        print(f"Status: {status}")
        print("-" * 50)

        if results["errors"]:
            print("\nErrors:")
            for e in results["errors"]:
                print(f"  ✗ {e}")

        if results["warnings"]:
            print("\nWarnings:")
            for w in results["warnings"]:
                print(f"  ⚠ {w}")

        if results["info"]:
            print("\nInfo:")
            for k, v in results["info"].items():
                print(f"  {k}: {v}")

    sys.exit(0 if results["valid"] else 1)


if __name__ == "__main__":
    main()
