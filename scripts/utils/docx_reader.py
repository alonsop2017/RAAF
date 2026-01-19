#!/usr/bin/env python3
"""
DOCX text extraction utility.
Uses python-docx for extracting text from Word documents.
"""

import sys
from pathlib import Path
from typing import Optional

try:
    from docx import Document
except ImportError:
    Document = None


def extract_text(docx_path: Path | str) -> str:
    """
    Extract text from a DOCX file.

    Args:
        docx_path: Path to the DOCX file

    Returns:
        Extracted text as a string

    Raises:
        FileNotFoundError: If the file doesn't exist
        ImportError: If python-docx is not installed
    """
    if Document is None:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    docx_path = Path(docx_path)

    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    doc = Document(docx_path)

    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))

    return "\n\n".join(text_parts)


def extract_text_with_metadata(docx_path: Path | str) -> dict:
    """
    Extract text and metadata from a DOCX file.

    Returns:
        Dictionary with 'text', 'metadata', 'source_file' keys
    """
    if Document is None:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    docx_path = Path(docx_path)
    doc = Document(docx_path)

    # Extract core properties
    props = doc.core_properties
    metadata = {
        "author": props.author,
        "created": str(props.created) if props.created else None,
        "modified": str(props.modified) if props.modified else None,
        "title": props.title,
        "subject": props.subject,
    }

    # Extract text
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))

    return {
        "text": "\n\n".join(text_parts),
        "metadata": metadata,
        "source_file": docx_path.name,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables)
    }


def extract_structured_content(docx_path: Path | str) -> dict:
    """
    Extract structured content from a DOCX file, preserving sections.

    Returns:
        Dictionary with sections based on heading styles
    """
    if Document is None:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    docx_path = Path(docx_path)
    doc = Document(docx_path)

    sections = []
    current_section = {"heading": None, "content": []}

    for para in doc.paragraphs:
        # Check if paragraph is a heading
        if para.style.name.startswith("Heading"):
            # Save previous section if it has content
            if current_section["content"] or current_section["heading"]:
                sections.append(current_section)
            current_section = {"heading": para.text, "content": []}
        elif para.text.strip():
            current_section["content"].append(para.text)

    # Don't forget the last section
    if current_section["content"] or current_section["heading"]:
        sections.append(current_section)

    return {
        "sections": sections,
        "source_file": docx_path.name
    }


def clean_extracted_text(text: str) -> str:
    """
    Clean up extracted text.

    Args:
        text: Raw extracted text

    Returns:
        Cleaned text
    """
    import re

    # Replace multiple spaces with single space
    text = re.sub(r' +', ' ', text)

    # Replace multiple newlines with double newline
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Remove leading/trailing whitespace from lines
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)

    return text.strip()


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python docx_reader.py <docx_file>")
        sys.exit(1)

    docx_file = Path(sys.argv[1])
    try:
        text = extract_text(docx_file)
        text = clean_extracted_text(text)
        print(text)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
